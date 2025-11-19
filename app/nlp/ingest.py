from __future__ import annotations

import os
import re
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Tuple

import pdfplumber
from docx import Document

SENT_SPLIT = re.compile(r"(?<=[\.?\!])\s+(?=[A-Z0-9])")


@dataclass
class Chunk:
    doc_id: str
    text: str
    meta: Dict[str, Any]


@dataclass
class StructuredSection:
    text: str
    meta: Dict[str, Any]


BAB_PATTERN = re.compile(r"^\s*BAB\s+([IVXLCDM]+|\d+)\s*(.*)$", re.IGNORECASE)
BAGIAN_PATTERN = re.compile(r"^\s*BAGIAN\s+([IVXLCDM]+|\d+)\s*(.*)$", re.IGNORECASE)
PARAGRAF_PATTERN = re.compile(r"^\s*PARAGRAF\s+([IVXLCDM]+|\d+)\s*(.*)$", re.IGNORECASE)
PASAL_PATTERN = re.compile(r"^\s*PASAL\s+(\d+[A-Z]?)\s*(.*)$", re.IGNORECASE)
AYAT_PATTERN = re.compile(r"^\s*(?:AYAT\s*)?\((\d+[A-Z]?)\)")


def _clean_text(value: str) -> str:
    value = value.replace("\u00ad", "")
    value = re.sub(r"\s+", " ", value).strip()
    return value


def read_pdf(path: str) -> List[Tuple[int, str]]:
    pages: List[Tuple[int, str]] = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            content = (page.extract_text() or "").replace("\u00ad", "")
            lines = []
            for raw_line in content.splitlines():
                cleaned_line = _clean_text(raw_line)
                if cleaned_line:
                    lines.append(cleaned_line)
            cleaned = "\n".join(lines)
            if cleaned:
                pages.append((idx, cleaned))
    return pages


def read_docx(path: str) -> List[Tuple[int, str]]:
    document = Document(path)
    blocks: List[Tuple[int, str]] = []
    buffer: List[str] = []
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            buffer.append(text)
        elif buffer:
            joined = "\n".join(buffer)
            blocks.append((len(blocks) + 1, joined))
            buffer = []
    if buffer:
        joined = "\n".join(buffer)
        blocks.append((len(blocks) + 1, joined))
    return blocks


def sentences(text: str) -> List[str]:
    return [sentence.strip() for sentence in SENT_SPLIT.split(text) if sentence.strip()]


def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> List[str]:
    chunks: List[str] = []
    current: List[str] = []
    size = 0
    for sentence in sentences(text):
        if size + len(sentence) > max_chars and current:
            chunk = " ".join(current)
            chunks.append(chunk)
            trailing = " ".join(chunk.split()[-overlap:]) if overlap else ""
            current = [trailing, sentence] if trailing else [sentence]
            size = len(" ".join(current))
        else:
            current.append(sentence)
            size += len(sentence)
    if current:
        chunks.append(" ".join(current))
    return chunks


class StructureTracker:
    """Tracks hierarchical legal structure (Bab, Pasal, Ayat, etc.) while parsing text."""

    def __init__(self) -> None:
        self.state: Dict[str, Any] = {
            "bab": None,
            "bab_title": None,
            "bagian": None,
            "bagian_title": None,
            "paragraf": None,
            "paragraf_title": None,
            "pasal": None,
            "ayat": None,
        }

    def snapshot(self) -> Dict[str, Any]:
        return {key: value for key, value in self.state.items() if value}

    def _reset_lower(self, *keys: str) -> None:
        for key in keys:
            if key in self.state:
                self.state[key] = None

    def apply_markers(self, markers: Dict[str, Any]) -> None:
        if "bab" in markers:
            self.state["bab"] = markers["bab"]
            if markers.get("bab_title"):
                self.state["bab_title"] = markers["bab_title"]
            self._reset_lower("bagian", "bagian_title", "paragraf", "paragraf_title", "pasal", "ayat")

        if "bagian" in markers:
            self.state["bagian"] = markers["bagian"]
            if markers.get("bagian_title"):
                self.state["bagian_title"] = markers["bagian_title"]
            self._reset_lower("paragraf", "paragraf_title", "pasal", "ayat")

        if "paragraf" in markers:
            self.state["paragraf"] = markers["paragraf"]
            if markers.get("paragraf_title"):
                self.state["paragraf_title"] = markers["paragraf_title"]
            self._reset_lower("pasal", "ayat")

        if "pasal" in markers:
            self.state["pasal"] = markers["pasal"]
            self._reset_lower("ayat")

        if "ayat" in markers:
            self.state["ayat"] = markers["ayat"]

    def detect_markers(self, text: str) -> Dict[str, Any]:
        markers: Dict[str, Any] = {}
        if match := BAB_PATTERN.match(text):
            markers["bab"] = match.group(1).strip()
            markers["bab_title"] = (match.group(2) or "").strip() or None
        elif match := BAGIAN_PATTERN.match(text):
            markers["bagian"] = match.group(1).strip()
            markers["bagian_title"] = (match.group(2) or "").strip() or None
        elif match := PARAGRAF_PATTERN.match(text):
            markers["paragraf"] = match.group(1).strip()
            markers["paragraf_title"] = (match.group(2) or "").strip() or None
        elif match := PASAL_PATTERN.match(text):
            markers["pasal"] = match.group(1).strip()
        elif match := AYAT_PATTERN.match(text):
            markers["ayat"] = match.group(1).strip()
        return markers


def _build_structured_sections(text: str, tracker: StructureTracker) -> List[StructuredSection]:
    sections: List[StructuredSection] = []
    buffer: List[str] = []
    normalized_text = text
    # ensure structure keywords begin on new lines even if original formatting was lost
    for keyword in ("BAB", "BAGIAN", "PARAGRAF", "PASAL", "AYAT"):
        normalized_text = re.sub(
            rf"(?i)(?<!\n)({keyword}\s+[IVXLCDM0-9]+)",
            r"\n\1",
            normalized_text,
        )
    normalized_text = re.sub(r"(?i)(?<!\n)((?:AYAT\s*)?\(\d+[A-Z]?\))", r"\n\1", normalized_text)

    for raw_line in normalized_text.splitlines():
        clean = _clean_text(raw_line)
        if not clean:
            continue
        markers = tracker.detect_markers(clean)
        if markers and buffer:
            sections.append(StructuredSection(text=" ".join(buffer), meta=tracker.snapshot()))
            buffer = []
        if markers:
            tracker.apply_markers(markers)
        buffer.append(clean)

    if buffer:
        sections.append(StructuredSection(text=" ".join(buffer), meta=tracker.snapshot()))

    if not sections and normalized_text.strip():
        sections.append(StructuredSection(text=_clean_text(normalized_text), meta=tracker.snapshot()))

    return sections


def _chunk_with_metadata(
    doc_id: str,
    source_label: str,
    origin: Iterable[Tuple[int, str]],
    origin_key: str,
    max_chars: int = 1200,
) -> List[Chunk]:
    chunks: List[Chunk] = []
    tracker = StructureTracker()
    section_counter = 0
    for origin_index, text in origin:
        structured_sections = _build_structured_sections(text, tracker)
        for section in structured_sections:
            if not section.text:
                continue
            section_counter += 1
            pieces = chunk_text(section.text, max_chars=max_chars)
            for part_index, piece in enumerate(pieces, start=1):
                meta = {
                    "source": source_label,
                    origin_key: origin_index,
                    "structured_section": section_counter,
                }
                meta.update(section.meta)
                if len(pieces) > 1:
                    meta["section_chunk"] = part_index
                chunks.append(Chunk(doc_id=doc_id, text=piece, meta=meta))
    return chunks


def build_chunks(path: str, doc_id: str) -> List[Chunk]:
    source_label = doc_id or os.path.basename(path)
    lower_path = path.lower()
    if lower_path.endswith(".pdf"):
        origin = read_pdf(path)
        return _chunk_with_metadata(doc_id, source_label, origin, "page")
    if lower_path.endswith(".docx"):
        origin = read_docx(path)
        if not origin:
            # Fallback: treat the whole document as one block if headings were empty.
            full_document = Document(path)
            cleaned_parts: List[str] = []
            for paragraph in full_document.paragraphs:
                cleaned = _clean_text(paragraph.text)
                if cleaned:
                    cleaned_parts.append(cleaned)
            combined = " ".join(cleaned_parts)
            origin = [(1, combined)] if combined else []
        return _chunk_with_metadata(doc_id, source_label, origin, "section")
    raise ValueError("Unsupported file type. Gunakan berkas .pdf atau .docx")
